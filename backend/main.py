from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from charset_normalizer import from_path
import shutil, os
import pandas as pd
from utils import call_llama3, classify_department, extract_entities, classify_ticket, generate_it_suggestion
from workflow_engine import generate_workflow
from pydantic import BaseModel
import bcrypt
import mysql.connector
from datetime import datetime

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # React frontend
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def get_db():
    return mysql.connector.connect(
        host="localhost",
        user="root",
        password="roylobo8",
        database="genai_portal"
    )

class User(BaseModel):
    username: str
    password: str
    department: str = "HR"

# ---------------- REGISTER & LOGIN ----------------
@app.post("/register")
def register(user: User):
    db = get_db()
    cursor = db.cursor()
    username = user.username.strip().lower()
    cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
    existing = cursor.fetchone()
    if existing:
        return {"success": False, "message": "User already exists"}
    hashed = bcrypt.hashpw(user.password.encode('utf-8'), bcrypt.gensalt())
    try:
        cursor.execute(
            "INSERT INTO users (username, password, department) VALUES (%s, %s, %s)",
            (username, hashed.decode('utf-8'), user.department)
        )
        db.commit()
        return {"success": True, "message": "User registered"}
    except mysql.connector.IntegrityError as e:
        if e.errno == 1062:
            return {"success": False, "message": "User already exists"}
        else:
            return {"success": False, "message": f"Database error: {str(e)}"}
    except Exception as e:
        return {"success": False, "message": f"Unexpected error: {str(e)}"}

@app.post("/login")
def login(user: User):
    db = get_db()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM users WHERE username=%s", (user.username,))
    record = cursor.fetchone()
    if record and bcrypt.checkpw(user.password.encode('utf-8'), record['password'].encode('utf-8')):
        return {"success": True, "user": {"id": record["id"], "username": record["username"], "department": record["department"]}}
    return {"success": False, "message": "Invalid credentials"}

# ---------------- IT TICKET SYSTEM ----------------
@app.post("/create_ticket/")
def create_ticket(user_id: int = Form(...), description: str = Form(...), affected_user: str = Form(None), ticket_type: str = Form("self")):
    db = get_db()
    cursor = db.cursor(dictionary=True)
    
    # Get user info for context
    cursor.execute("SELECT username, department FROM users WHERE id=%s", (user_id,))
    user_info = cursor.fetchone()
    
    # Determine ticket context
    if ticket_type == "self":
        ticket_context = f"Self-reported by {user_info['username']}"
    elif ticket_type == "other" and affected_user:
        ticket_context = f"Reported by {user_info['username']} for {affected_user}"
    elif ticket_type == "system":
        ticket_context = f"System-generated ticket for {affected_user or 'unknown user'}"
    else:
        ticket_context = f"Reported by {user_info['username']}"
    
    # Add context to description
    full_description = f"{ticket_context}\n\nIssue: {description}"
    
    category = classify_ticket(description)
    ai_summary = call_llama3(f"Summarize this IT ticket: {full_description[:2000]}")
    ai_suggestion = generate_it_suggestion(category, description)
    
    try:
        cursor.execute("""
            INSERT INTO tickets (user_id, category, description, ai_summary, ai_suggestion, status, affected_user, ticket_type)
            VALUES (%s, %s, %s, %s, %s, 'Open', %s, %s)
        """, (user_id, category, full_description, ai_summary, ai_suggestion, affected_user, ticket_type))
    except Exception as e:
        # If new columns don't exist, try without them
        print(f"Error with new columns, trying without: {str(e)}")
        try:
            cursor.execute("""
                INSERT INTO tickets (user_id, category, description, ai_summary, ai_suggestion, status)
                VALUES (%s, %s, %s, %s, %s, 'Open')
            """, (user_id, category, full_description, ai_summary, ai_suggestion))
        except Exception as e2:
            # If status column doesn't exist, try without it
            print(f"Error with status column, trying without: {str(e2)}")
            cursor.execute("""
                INSERT INTO tickets (user_id, category, description, ai_summary, ai_suggestion)
                VALUES (%s, %s, %s, %s, %s)
            """, (user_id, category, full_description, ai_summary, ai_suggestion))
    
    db.commit()
    
    return {
        "success": True,
        "category": category,
        "summary": ai_summary,
        "suggestion": ai_suggestion,
        "context": ticket_context
    }

@app.get("/tickets/")
def get_tickets(user_id: int):
    db = get_db()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM tickets WHERE user_id=%s ORDER BY created_at DESC", (user_id,))
    return cursor.fetchall()

@app.post("/query_ticket/")
def query_ticket(user_id: int = Form(...), question: str = Form(...)):
    db = get_db()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT ai_summary FROM tickets WHERE user_id=%s", (user_id,))
    summaries = " ".join([t["ai_summary"] for t in cursor.fetchall()])
    prompt = f"User submitted IT tickets summaries:\n{summaries}\nQuestion: {question}"
    answer = call_llama3(prompt)
    return {"answer": answer}

@app.get("/download_csv/")
def download_csv(user_id: int):
    db = get_db()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM tickets WHERE user_id=%s", (user_id,))
    tickets = cursor.fetchall()
    if not tickets:
        return {"error": "No tickets"}
    df = pd.DataFrame(tickets)
    path = f"user_{user_id}_tickets.csv"
    df.to_csv(path, index=False)
    return FileResponse(path, filename=path)

@app.post("/resolve_ticket/")
def resolve_ticket(ticket_id: int = Form(...), user_id: int = Form(...)):
    try:
        print(f"Attempting to resolve ticket {ticket_id} for user {user_id}")
        
        # Test database connection first
        try:
            db = get_db()
            cursor = db.cursor(dictionary=True)
            print("Database connection successful")
        except Exception as conn_error:
            print(f"Database connection failed: {str(conn_error)}")
            return {"success": False, "message": f"Database connection error: {str(conn_error)}"}
        
        # Try to add status column if it doesn't exist (ignore if already exists)
        try:
            cursor.execute("ALTER TABLE tickets ADD COLUMN status VARCHAR(20) DEFAULT 'Open'")
            db.commit()
            print("Status column added successfully")
        except Exception as alter_error:
            # If column already exists, this will fail - that's okay
            print(f"Status column already exists or other alter error: {str(alter_error)}")
            db.rollback()  # Rollback the failed alter
        
        # Verify the ticket belongs to the user and check if they can resolve it
        try:
            cursor.execute("SELECT * FROM tickets WHERE id=%s AND user_id=%s", (ticket_id, user_id))
            ticket = cursor.fetchone()
            print(f"Found ticket: {ticket}")
            
            if not ticket:
                return {"success": False, "message": "Ticket not found or unauthorized"}
            
            # Check if user can resolve this ticket
            # Get user info to check username
            cursor.execute("SELECT username FROM users WHERE id=%s", (user_id,))
            user_info = cursor.fetchone()
            
            # User can resolve if:
            # 1. It's their own ticket (ticket_type = "self" or no affected_user)
            # 2. They are the affected user
            # 3. No affected_user specified (legacy tickets)
            can_resolve = (
                ticket.get('ticket_type') == 'self' or
                not ticket.get('affected_user') or
                ticket.get('affected_user') == user_info['username']
            )
            
            if not can_resolve:
                return {"success": False, "message": "Only the affected user can resolve this ticket"}
                
        except Exception as select_error:
            print(f"Error selecting ticket: {str(select_error)}")
            return {"success": False, "message": f"Error finding ticket: {str(select_error)}"}
        
        # Update ticket status to resolved
        try:
            cursor.execute("UPDATE tickets SET status='Resolved' WHERE id=%s", (ticket_id,))
            db.commit()
            print(f"Successfully updated ticket {ticket_id} to resolved")
        except Exception as update_error:
            print(f"Error updating ticket: {str(update_error)}")
            return {"success": False, "message": f"Error updating ticket: {str(update_error)}"}
        
        return {"success": True, "message": "Ticket marked as resolved"}
    except Exception as e:
        print(f"General error resolving ticket: {str(e)}")
        return {"success": False, "message": f"Database error: {str(e)}"}

@app.post("/reopen_ticket/")
def reopen_ticket(ticket_id: int = Form(...), user_id: int = Form(...), reason: str = Form(...)):
    try:
        print(f"Attempting to reopen ticket {ticket_id} for user {user_id}")
        
        # Test database connection first
        try:
            db = get_db()
            cursor = db.cursor(dictionary=True)
            print("Database connection successful")
        except Exception as conn_error:
            print(f"Database connection failed: {str(conn_error)}")
            return {"success": False, "message": f"Database connection error: {str(conn_error)}"}
        
        # Try to add status column if it doesn't exist (ignore if already exists)
        try:
            cursor.execute("ALTER TABLE tickets ADD COLUMN status VARCHAR(20) DEFAULT 'Open'")
            db.commit()
            print("Status column added successfully")
        except Exception as alter_error:
            # If column already exists, this will fail - that's okay
            print(f"Status column already exists or other alter error: {str(alter_error)}")
            db.rollback()  # Rollback the failed alter
        
        # Verify the ticket belongs to the user
        try:
            cursor.execute("SELECT * FROM tickets WHERE id=%s AND user_id=%s", (ticket_id, user_id))
            ticket = cursor.fetchone()
            print(f"Found ticket: {ticket}")
            
            if not ticket:
                return {"success": False, "message": "Ticket not found or unauthorized"}
        except Exception as select_error:
            print(f"Error selecting ticket: {str(select_error)}")
            return {"success": False, "message": f"Error finding ticket: {str(select_error)}"}
        
        # Update ticket status to reopened and add escalation reason
        try:
            cursor.execute("UPDATE tickets SET status='Reopened', escalation_reason=%s WHERE id=%s", (reason, ticket_id))
            db.commit()
            print(f"Successfully reopened ticket {ticket_id}")
        except Exception as update_error:
            print(f"Error reopening ticket: {str(update_error)}")
            return {"success": False, "message": f"Error reopening ticket: {str(update_error)}"}
        
        return {"success": True, "message": "Ticket reopened successfully"}
    except Exception as e:
        print(f"General error reopening ticket: {str(e)}")
        return {"success": False, "message": f"Database error: {str(e)}"}

@app.post("/escalate_ticket/")
def escalate_ticket(ticket_id: int = Form(...), user_id: int = Form(...), escalation_reason: str = Form(...)):
    try:
        print(f"Attempting to escalate ticket {ticket_id} for user {user_id}")
        
        # Test database connection first
        try:
            db = get_db()
            cursor = db.cursor(dictionary=True)
            print("Database connection successful")
        except Exception as conn_error:
            print(f"Database connection failed: {str(conn_error)}")
            return {"success": False, "message": f"Database connection error: {str(conn_error)}"}
        
        # Try to add escalation_reason column if it doesn't exist
        try:
            cursor.execute("ALTER TABLE tickets ADD COLUMN escalation_reason TEXT")
            db.commit()
            print("Escalation reason column added successfully")
        except Exception as alter_error:
            print(f"Escalation reason column already exists or other alter error: {str(alter_error)}")
            db.rollback()
        
        # Verify the ticket belongs to the user
        try:
            cursor.execute("SELECT * FROM tickets WHERE id=%s AND user_id=%s", (ticket_id, user_id))
            ticket = cursor.fetchone()
            print(f"Found ticket: {ticket}")
            
            if not ticket:
                return {"success": False, "message": "Ticket not found or unauthorized"}
        except Exception as select_error:
            print(f"Error selecting ticket: {str(select_error)}")
            return {"success": False, "message": f"Error finding ticket: {str(select_error)}"}
        
        # Update ticket status to escalated
        try:
            cursor.execute("UPDATE tickets SET status='Escalated', escalation_reason=%s WHERE id=%s", (escalation_reason, ticket_id))
            db.commit()
            print(f"Successfully escalated ticket {ticket_id}")
        except Exception as update_error:
            print(f"Error escalating ticket: {str(update_error)}")
            return {"success": False, "message": f"Error escalating ticket: {str(update_error)}"}
        
        return {"success": True, "message": "Ticket escalated successfully"}
    except Exception as e:
        print(f"General error escalating ticket: {str(e)}")
        return {"success": False, "message": f"Database error: {str(e)}"}

@app.get("/debug/tickets/")
def debug_tickets():
    """Debug endpoint to check ticket table structure"""
    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        
        # Check if table exists by trying to select from it
        cursor.execute("SELECT COUNT(*) as count FROM tickets")
        count_result = cursor.fetchone()
        
        # Check sample data
        cursor.execute("SELECT * FROM tickets LIMIT 3")
        sample_tickets = cursor.fetchall()
        
        return {
            "table_exists": True,
            "total_tickets": count_result['count'] if count_result else 0,
            "sample_tickets": sample_tickets
        }
    except Exception as e:
        return {"error": str(e), "table_exists": False}

# ---------------- HR TALENT MANAGEMENT SYSTEM ----------------
@app.post("/upload_resume/")
async def upload_resume(file: UploadFile = File(...), user_id: int = Form(...)):
    """Upload and analyze resume for job matching"""
    try:
        print(f"Starting resume upload for user {user_id}, file: {file.filename}")
        
        # Ensure upload directory exists
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        
        file_path = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        print(f"File saved to: {file_path}")
    
        text = ""
        filename_lower = file.filename.lower().strip()
        try:
            if filename_lower.endswith(".txt"):
                detected = from_path(file_path).best()
                text = str(detected) if detected else ""
            elif filename_lower.endswith(".pdf"):
                import pdfplumber
                try:
                    with pdfplumber.open(file_path) as pdf:
                        if len(pdf.pages) == 0:
                            return JSONResponse(content={"error": f"PDF {file.filename} has no readable pages"})
                        text = " ".join([page.extract_text() or "" for page in pdf.pages])
                        if not text.strip():
                            return JSONResponse(content={"error": f"PDF {file.filename} contains no extractable text"})
                except Exception as pdf_error:
                    return JSONResponse(content={"error": f"PDF processing failed for {file.filename}", "details": str(pdf_error)})
            elif filename_lower.endswith(".docx"):
                from docx import Document
                doc = Document(file_path)
                text = " ".join([p.text for p in doc.paragraphs])
        except Exception as e:
            return JSONResponse(content={"error": f"Failed to read {file.filename}", "details": str(e)})
        
        if not text.strip():
            return JSONResponse(content={"error": "No readable text found", "filename": file.filename})
        
        print(f"Extracted text length: {len(text)} characters")
        
        # Analyze resume and extract skills
        try:
            skills_prompt = f"""
            Analyze this resume and extract:
            1. Candidate name
            2. Years of experience
            3. Technical skills (programming languages, tools, frameworks)
            4. Soft skills (communication, leadership, etc.)
            5. Education background
            6. Previous job roles
            
            Resume: {text[:3000]}
            
            Return as JSON format:
            {{
                "name": "candidate name",
                "experience_years": number,
                "technical_skills": ["skill1", "skill2"],
                "soft_skills": ["skill1", "skill2"],
                "education": "degree and institution",
                "previous_roles": ["role1", "role2"]
            }}
            """
            
            print("Calling LLM for skills analysis...")
            skills_analysis = call_llama3(skills_prompt)
            print(f"Skills analysis completed: {len(skills_analysis)} characters")
        except Exception as e:
            print(f"Error in skills analysis: {str(e)}")
            return JSONResponse(content={"error": f"Skills analysis failed: {str(e)}"})
        
        # Job matching analysis
        try:
            job_matching_prompt = f"""
            Based on this candidate profile, analyze their fit for different job roles:
            
            {skills_analysis}
            
            Available job roles:
            1. Frontend Developer (React, Vue, Angular, JavaScript, HTML, CSS)
            2. Backend Developer (Python, Java, Node.js, SQL, APIs)
            3. Full Stack Developer (Frontend + Backend skills)
            4. Data Analyst (SQL, Python, Excel, Tableau, PowerBI)
            5. DevOps Engineer (Docker, Kubernetes, AWS, CI/CD)
            6. UI/UX Designer (Figma, Adobe, User Research, Prototyping)
            7. Project Manager (Agile, Scrum, Leadership, Communication)
            8. Business Analyst (Requirements, Documentation, Stakeholder Management)
            9. QA Engineer (Testing, Automation, Selenium, JUnit)
            10. Support Engineer (Customer Service, Technical Support, Troubleshooting)
            11. Sales Executive (Sales, CRM, Communication, Negotiation)
            12. Marketing Specialist (Digital Marketing, SEO, Social Media, Analytics)
            13. Finance Analyst (Accounting, Excel, Financial Modeling, Analysis)
            14. HR Specialist (Recruitment, Employee Relations, HRIS, Compliance)
            15. Operations Manager (Process Improvement, Team Management, Logistics)
            
            Return ONLY a JSON array with this exact format:
            [
              {{
                "role": "Role Name",
                "match": 85,
                "fit": "High"
              }},
              {{
                "role": "Role Name", 
                "match": 72,
                "fit": "Medium"
              }}
            ]
            
            Rules:
            - Return maximum 3 best-fit roles
            - Match percentage should be 0-100
            - Fit should be "High" (80+), "Medium" (60-79), or "Low" (below 60)
            - Only return the JSON array, no other text
            """
            
            print("Calling LLM for job matching...")
            job_matches = call_llama3(job_matching_prompt)
            print(f"Job matching completed: {len(job_matches)} characters")
        except Exception as e:
            print(f"Error in job matching: {str(e)}")
            return JSONResponse(content={"error": f"Job matching failed: {str(e)}"})
        
        # Store in database
        try:
            print("Connecting to database...")
            db = get_db()
            cursor = db.cursor(dictionary=True)
            
            try:
                cursor.execute("""
                    INSERT INTO resumes (user_id, filename, candidate_name, experience_years, 
                                       technical_skills, soft_skills, education, previous_roles, 
                                       skills_analysis, job_matches, status)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 'Analyzed')
                """, (user_id, file.filename, "", 0, "", "", "", "", skills_analysis, job_matches))
                print("Resume inserted successfully")
            except Exception as e:
                # If table doesn't exist, create it
                print(f"Error inserting resume: {str(e)}")
                print("Creating resumes table...")
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS resumes (
                        id INT AUTO_INCREMENT PRIMARY KEY,
                        user_id INT,
                        filename VARCHAR(255),
                        candidate_name VARCHAR(255),
                        experience_years INT,
                        technical_skills TEXT,
                        soft_skills TEXT,
                        education TEXT,
                        previous_roles TEXT,
                        skills_analysis TEXT,
                        job_matches TEXT,
                        status VARCHAR(50),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                """)
                db.commit()
                print("Resumes table created successfully")
                
                # Try insert again
                cursor.execute("""
                    INSERT INTO resumes (user_id, filename, candidate_name, experience_years, 
                                       technical_skills, soft_skills, education, previous_roles, 
                                       skills_analysis, job_matches, status)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, 'Analyzed')
                """, (user_id, file.filename, "", 0, "", "", "", "", skills_analysis, job_matches))
                print("Resume inserted after table creation")
            
            db.commit()
            print("Database commit successful")
        except Exception as e:
            print(f"Database error: {str(e)}")
            return JSONResponse(content={"error": f"Database error: {str(e)}"})
        
        return {
            "success": True,
            "filename": file.filename,
            "skills_analysis": skills_analysis,
            "job_matches": job_matches
        }
    except Exception as e:
        print(f"General error in resume upload: {str(e)}")
        return JSONResponse(content={"error": f"Resume upload failed: {str(e)}"})

@app.get("/resumes/")
def get_resumes(user_id: int):
    """Get all resumes for a user"""
    db = get_db()
    cursor = db.cursor(dictionary=True)
    
    try:
        cursor.execute("SELECT * FROM resumes WHERE user_id=%s ORDER BY created_at DESC", (user_id,))
        return cursor.fetchall()
    except Exception as e:
        # If table doesn't exist, return empty
        return []

@app.post("/search_candidates/")
def search_candidates(user_id: int = Form(...), job_role: str = Form(...), min_experience: int = Form(0)):
    """Search for candidates matching specific job criteria"""
    db = get_db()
    cursor = db.cursor(dictionary=True)
    
    try:
        cursor.execute("""
            SELECT * FROM resumes 
            WHERE user_id=%s AND job_matches LIKE %s AND experience_years >= %s
            ORDER BY created_at DESC
        """, (user_id, f"%{job_role}%", min_experience))
        return cursor.fetchall()
    except Exception as e:
        return []

# ---------------- EXISTING DOCUMENT WORKFLOW ----------------
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
results_store = []

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    text = ""
    filename_lower = file.filename.lower().strip()
    try:
        if filename_lower.endswith(".txt"):
            detected = from_path(file_path).best()
            text = str(detected) if detected else ""
        elif filename_lower.endswith(".pdf"):
            import pdfplumber
            try:
                with pdfplumber.open(file_path) as pdf:
                    if len(pdf.pages) == 0:
                        return JSONResponse(content={"error": f"PDF {file.filename} has no readable pages"})
                    text = " ".join([page.extract_text() or "" for page in pdf.pages])
                    if not text.strip():
                        return JSONResponse(content={"error": f"PDF {file.filename} contains no extractable text"})
            except Exception as pdf_error:
                return JSONResponse(content={"error": f"PDF processing failed for {file.filename}", "details": str(pdf_error)})
        elif filename_lower.endswith(".docx"):
            from docx import Document
            doc = Document(file_path)
            text = " ".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return JSONResponse(content={"error": f"Failed to read {file.filename}", "details": str(e)})
    
    if not text.strip():
        return JSONResponse(content={"error": "No readable text found", "filename": file.filename})
    
    summary = call_llama3(f"Summarize this document:\n{text[:2000]}")
    dept = classify_department(text)
    entities = extract_entities(text)
    
    # Add raw text to entities for workflow analysis
    entities["raw"] = text
    entities["summary"] = summary
    
    workflow = generate_workflow(dept, entities)
    
    result = {
        "filename": file.filename,
        "department": dept,
        "summary": summary,
        "entities": entities,
        "workflow_outcome": workflow["outcome"],
        "workflow_checklist": workflow["checklist"]
    }
    
    results_store.append(result)
    return JSONResponse(content=result)

@app.get("/results/")
async def get_results():
    # Filter out entries whose source files have been removed from UPLOAD_DIR
    existing = []
    for r in results_store:
        try:
            path = os.path.join(UPLOAD_DIR, r.get("filename", ""))
            if path and os.path.exists(path):
                existing.append(r)
        except Exception:
            # If any issue occurs while checking, skip that entry
            pass

    # Keep in-memory store consistent if anything was removed
    if len(existing) != len(results_store):
        results_store[:] = existing

    return existing

@app.post("/query/")
async def query_docs(question: str = Form(...)):
    all_texts = " ".join([r["summary"] for r in results_store])
    answer = call_llama3(f"Answer this based on docs:\n{all_texts}\nQuestion: {question}")
    return {"answer": answer}

@app.get("/download_docs/")
async def download_docs_csv():
    if not results_store:
        return {"error": "No data available"}
    df = pd.DataFrame(results_store)
    csv_path = "results_docs.csv"
    df.to_csv(csv_path, index=False)
    return FileResponse(csv_path, filename="results_docs.csv")
